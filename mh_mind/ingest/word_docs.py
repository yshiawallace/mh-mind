"""Ingest Word documents from configured folder paths.

Walks each folder recursively, parses .docx files with python-docx,
and returns WordDoc dataclasses. Skips .doc (pre-2007) files with a warning.
"""

import logging
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from mh_mind.ingest.footnotes import paragraph_text_with_notes, parse_notes

logger = logging.getLogger(__name__)


@dataclass
class WordDoc:
    title: str
    body: str
    path: Path
    modified: datetime


def _parse_docx(path: Path) -> WordDoc | None:
    """Parse a single .docx file into a WordDoc."""
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("Could not parse %s: %s", path, e)
        return None

    # Parse footnotes and endnotes from the .docx XML
    notes = parse_notes(path)
    if notes:
        logger.info("Found %d footnotes/endnotes in %s", len(notes), path.name)

    # Extract text from paragraphs (with inlined notes) and table cells
    parts: list[str] = []

    for para in doc.paragraphs:
        text = paragraph_text_with_notes(para._element, notes)
        if text.strip():
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    body = "\n\n".join(parts)

    # Use filename (without extension) as title
    title = path.stem

    # Use filesystem modification time
    modified = datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc)

    return WordDoc(title=title, body=body, path=path, modified=modified)


def load_docs(folder_paths: list[Path]) -> list[WordDoc]:
    """Walk the given folders and parse all .docx files found.

    Args:
        folder_paths: Directories to search recursively for .docx files.

    Returns:
        List of parsed WordDoc dataclasses.
    """
    docs: list[WordDoc] = []

    for folder in folder_paths:
        if not folder.is_dir():
            logger.warning("Skipping non-existent folder: %s", folder)
            continue

        for path in sorted(folder.rglob("*")):
            # Skip hidden files and directories
            if any(part.startswith(".") for part in path.parts):
                continue

            if path.suffix.lower() == ".doc":
                logger.warning(
                    "Skipping pre-2007 .doc file (not supported by python-docx): %s", path
                )
                continue

            if path.suffix.lower() != ".docx":
                continue

            # Skip temp files created by Word (e.g. ~$document.docx)
            if path.name.startswith("~$"):
                continue

            doc = _parse_docx(path)
            if doc is not None:
                docs.append(doc)

    logger.info("Loaded %d Word documents from %d folders", len(docs), len(folder_paths))
    return docs
